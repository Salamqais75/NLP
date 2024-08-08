from docx import Document
import zipfile
import os, re, json, sys, pandas


def get_docx(folder_path):
    try:

        info = []
        # go thru all the documants in the directory
        current_path = os.path.join(os.getcwd(), folder_path)

        for _, filename in enumerate(os.listdir(current_path)):
            if filename.endswith('docx'):

                attributes = filename.split('_')  # Split the filename to get the attributes

                if attributes[1] == 'ptv':  # committee or plenary
                    type = 'committee'
                elif attributes[1] == 'ptm':
                    type = 'plenary'
                else:
                    type = '-1'

                text = Document(os.path.join(current_path, filename))

                info_item = {
                    'file_name': filename,
                    'knesset number': int(attributes[0]),
                    'type': type,
                    'text': text,
                    'file_number': attributes[2].replace('.docx', '')
                }
                info.append(info_item)
        return info
    except Exception as e:
        print(f'Exception in get_docx: {e}')


# Function to the number after we find "הישיבה ה" or "'פרוטוקול מס"
def get_next_word(text, position):
    # Find the start of the next word
    word_start = position

    while word_start < len(text) and text[word_start].isspace():  # Skip spaces
        word_start += 1

    # If reached the end of the text, return -1
    if word_start >= len(text):
        return '-1'

    # Find the end of the next word
    word_end = word_start
    while word_end < len(text) and not text[word_end].isspace():
        word_end += 1

    # Return the next continuous word
    return text[word_start:word_end]


def is_underlined(par):
    try:
        # Check if all text runs are underlined
        all_underlined = True  # Assume all are underlined unless proven otherwise
        for run in par.runs:
            if not run.underline:
                all_underlined = False  # If any run is not underlined, then it's not fully underlined

        if all_underlined:
            return True

        # If any run is underlined, consider it partially underlined
        for run in par.runs:
            if run.underline:
                return True

        # Check the paragraph style for underline attributes
        current_style = par.style
        while current_style:
            if current_style.font.underline:
                return True  # If underline is defined at the style level

            # Check base styles for inherited underline
            current_style = current_style.base_style

        return False  # Default to not underlined
    except Exception as e:
        print(f"Exception in is_underlined: {e}")


def clean_name(name):  # clean
    try:
        name = name.strip()
        comps = name.split(' ')  # split name to words/comps
        clean_name = ""
        open_parentheses = False
        common_pos = ["ראש", "ממשלה", "יו\"ר", '”', "לאומי", "ערבית", "ועדת", "איכות", "פנים", "שר", "אוצר", "משפטים",
                      "\"", "אנרגיה", "מים", "תשתיות"]  # List of all the common positions

        for comp in comps:
            if comp == '':
                continue
            if any(pos in comp for pos in common_pos):  # if the text contains any of the common positions then skip
                continue
            if '(' in comp:  # This is a closing parentheses, backwards since its in hebrew
                open_parentheses = False
                continue
            if open_parentheses:
                continue
            if '\'' == comp[-1] and len(comp) < 4:  # then this a position
                clean_name = ''  # if it is then remove what we collected before
                continue

            if ")" in comp:  # This is an opening parentheses, backwads since its in hebrew
                if "(" in comp:  # if we have closing parentheses then skip it
                    continue
                else:  # else we must take the following comps until we see closing parentheses
                    open_parentheses = True

            elif comp == "-" or comp == '–' or comp == '~' or comp == ',':  # if the name has a dash then take the first part
                break
            else:
                clean_name += comp + " "

        clean_name = clean_name.strip()
        if ',' in clean_name:
            return ''
        if clean_name != "" and clean_name.find(':') + 1 == len(
                clean_name):  # Remove the colon if it's the last character, check if name isnt empty
            clean_name = clean_name[:-1]

        return clean_name.strip()
    except Exception as e:
        print(f'Exception in clear_name: {e}')


def split_paragrph(par):
    try:
        new_sentence = ''
        seperators = '.؟!?!;:'  # use this to check start and end of sentences
        qutoed = False

        cases = [' - - -','- - -', ' - -' , '- -' , ' – – –','– – –', ' – –' , '– –' , ' – – –', '– – –', ' – –' , '– –' ]
        txt = par.text.strip()
        for case in cases:
            if case in txt:  # Check if the text contains any of the special cases
                txt = txt.replace(txt, '')  # Replace the special case with a space
                break

        par_parts = txt.split(' ')
        sentece_list = []
        for part in par_parts:
            if part == '':  # if empty then skip
                continue

            new_sentence += part + " "  # Collect sentence
            if '"' == part[0]:
                qutoed = True
            if '"' == part[-1] or (len(part) >= 2 and part[-2] == '"' and part[-1] in seperators):
                # if the second to last char is a " and last is a seperator, or if the last is a quote then we end the quote
                qutoed = False

            if part[-1] in seperators or (
                    len(part) >= 2 and part[-2] in seperators):  # If we reached the end of the sentence, save it
                if qutoed == False:  # if were still in quotes, we dont save yet
                    sentece_list.append(new_sentence.strip())
                    new_sentence = ''

        # if we start a quote but it didnt end, save the text
        if qutoed:
            sentece_list.append(new_sentence)
        return sentece_list
    except Exception as e:
        print(f'exception in split_paragrph: {e}')


def remove_tags(text, tags):
    # Strip leading/trailing spaces
    try:
        cleaned_text = text.strip()
        for tag in tags:
            # Remove specific tag from start
            if cleaned_text.startswith(tag):
                cleaned_text = cleaned_text[len(tag):].strip()  # Remove the tag and strip spaces

            # Remove specific tag from end
            if cleaned_text.endswith(tag):
                cleaned_text = cleaned_text[:-len(tag)].strip()  # Remove the tag and strip spaces

        return cleaned_text
    except Exception as e:
        print(f'Exception in remove_tags: {e}')


def clean_text(txt):
    try:
        if txt == '':
            return ''
        allowed = re.compile('[א-ת0-9!"#$%&\'()*+,-./:;<=>?@[\\]_`{|}~– ]+')  # Allowed characters
        occurences = re.findall(allowed, txt)  # Find all allowed characters
        if len(occurences) != 1:  # If there are more than one occurence that means we have unwated characters in the text or in between the text
            return ''

        filtered_txt = occurences[0]  # Get the first and only occurence, we must check if its actually hebrew or not
        heb_txt = False  # Check if the text is in hebrew
        heb_letters = ['א', 'ב', 'ג', 'ד', 'ה', 'ו', 'ז', 'ח', 'ט', 'י', 'ך', 'כ', 'ל', 'ם', 'מ', 'ן', 'נ', 'ס', 'ע',
                       'ף', 'פ', 'ץ', 'צ', 'ק', 'ר', 'ש', 'ת']

        for letter in filtered_txt:
            if letter in heb_letters:
                heb_txt = True  # This means that the text is in hebrew since theres one occurence
                break
        if heb_txt == False:
            return ''

        return txt
    except Exception as e:
        print(f'Exception in clean_text: {e}')




def tokenize(list_text):
    try:
        tokens = []
        punctuation = ['!', '"', '\'', '(', ')', ',', '-', '.', '/', ':', ';', '?', '[', '\\', ']', '_', '{', '}', '~']

        for text in list_text:
            words = text.split(' ')  # split the text into words
            new_token = []
            for j in range(len(words)):
                only_punctuation = True  # Check if the word is only punctuation
                word = words[j]
                if word == '':
                    continue
                for i in range(len(word)):
                    if only_punctuation == False:  # this is done so we save text like 3:00 without spearing the colon
                        break
                    if word[i] in punctuation:
                        new_token.append(word[i])
                        words[j] = words[j][1:]  # Remove from the text
                    else:
                        only_punctuation = False

                punctuation_at_end = []  # save extra seperated punctuation marks at the end of the word
                for i in reversed(range(len(word))):
                    if word[i] in punctuation:
                        punctuation_at_end.append(word[i])
                        words[j] = words[j][:-1]  # Remove from the text
                    else:
                        new_token.append(words[j])
                        new_token.extend(reversed(punctuation_at_end))
                        break

            if len(new_token) < 4:
                continue
            tokens.append(new_token)
        return tokens

    except Exception as e:
        print(f'Exception in tokenize: {e}')


def fix_protocol(str):
    # Regular expressions for numbers only and letters only

    digit_pattern = r'^\d+$'
    first = {"עשרי": 20, "עשר": 10, "שלושים": 30, "ארבעים": 40, "חמישים": 50, "שישים": 60, "שבעים": 70, "שמונים": 80,
             "תשעים": 90}
    second = {"אח": 1, "שתי": 2, "שלוש": 3, "ארבע": 4, "חמש": 5, "שש": 6, "שבע": 7, "שמונה": 8, "תשע": 9}

    # Check if the input string matches the digit pattern
    if re.match(digit_pattern, str):
        return int(str)
    num = 0
    splits = str.split('-')
    for i in range(len(splits)):
        enter = True

        if splits[i] == '':
            continue
        if "מאה" in splits[i]:
            num += 100
            continue
        if "מאתי" in splits[i]:
            num += 200
            continue
        if "מאות" in splits[i]:
            num *= 100
            continue

        for key in first:
            if key in splits[i]:
                num += first[key]
                enter = False
                break

        if enter == False:
            continue
        for key in second:
            if key in splits[i]:
                num += second[key]
    return num


if __name__ == "__main__":
    try:
        # print(sys.argv)
        if len(sys.argv) != 3:
            print('Incorrect input, please enter the folder path and the output path.')
            sys.exit(1)

        # Add Check if folder is valid

        # CHECK IF ITS IN CORRECT JSONL FORMAT
        folder_path = sys.argv[1]
        output_path = sys.argv[2]
        info = get_docx(folder_path)

        tags = ["<< דובר >>", "<< נושא >>", "<< יור >>", "<< דובר_המשך >>", "<< אורח >>", "<< סיום >>", "<< הפסקה >>",
                "<< יור >>"]  # Tags to remove from the text
        common_pos = ["סדר-היום", "סדר היום", "נכחו", "חברי", "מנהל", "רישום", "משתתפים", "מוזמנים", "ייעוץ", "יועץ",
                      "קצרנית", "יועצת", "קצרן"]  # List of all the common positions
        target_words = ["הישיבה ה", "פרוטוקול מס'"]  # Search for the protocal number
        jsonl_data = []
        names = []

        # print(split_paragrph("asd ,dsa, 1:2:3:"))

        for doc_num, doc in enumerate(info):
            knesset_number = doc['knesset number']
            protocol_type = doc['type']
            file_number = doc['file_number']
            protocol_name = doc['file_name']
            protocol_number = '-1'  # Default value

            speakers_order = []  # list to hold the order of the speakers

            # Find the protocol number
            for par in doc['text'].paragraphs:
                text = par.text.strip()  # Remove leading and trailing spaces

                if text.startswith('<') or text.startswith(
                        '>'):  # Sometimes the text starts with < and ends with >, its probably caused by the conversion
                    text = text[1:-1]

                if target_words[0] in text:
                    position = text.find(target_words[0])
                    next_word = get_next_word(text, position + len(target_words[0]))
                    # print(f"Found in doc {doc_num}: {text}, NEXT {next_word}") # Debugging
                    protocol_number = next_word
                    break

                if target_words[1] in text:
                    position = text.find(target_words[1])
                    next_word = get_next_word(text, position + len(target_words[1]))
                    # print(f"Found in doc {doc_num}: {text}, NEXT {next_word}") # Debugging
                    protocol_number = next_word
                    break

            if protocol_number != '-1':
                if protocol_number[-1] == ',' or protocol_number[
                    -1] == '.':  # Remove the last character if it's a comma or period
                    protocol_number = protocol_number[:-1]
                protocol_int = fix_protocol(protocol_number)
                # print(f"Protocol number: {protocol_int}, string {protocol_number}")
            else:
                protocol_int = -1
            # Extract speakers and text

            prev_speaker = ''  # name of the first speaker
            speaker_text = {}  # dic to hold the text of each speaker

            for par in doc['text'].paragraphs:
                text = par.text
                text = remove_tags(text, tags)  # remove text tags
                if text.startswith('<') or text.startswith(
                        '>'):  # Sometimes the text starts with < and ends with >, its probably caused by the conversion
                    text = text[1:-1]

                    # if the last char is : and the whole text is underlined then this is a speaker
                index = text.find(":")
                if index >= 0 and index == len(text) - 1 and is_underlined(par):
                    if any(pos in text for pos in
                           common_pos):  # if the text contains any of the common positions then skip
                        continue
                    new_name = clean_name(text)  # clean the name

                    if new_name != '':
                        prev_speaker = new_name

                    else:  # if the name is empty then skip
                        split_txt = split_paragrph(par)

                        if prev_speaker != '':
                            filtered_text = []
                            for sent in split_txt:
                                filtered = clean_text(sent)
                                if filtered != '':
                                    filtered_text.append(filtered)

                            all_tokens = tokenize(filtered_text)
                            if len(all_tokens) == 0:
                                continue

                            for token in all_tokens:
                                combine_tokens = ''
                                for word1 in token:
                                    combine_tokens += str(word1) + ' '
                                speaker_text[prev_speaker].append(combine_tokens.strip())

                    if prev_speaker not in list(speaker_text.keys()):  # if the speaker is not in the dict then add him
                        speaker_text[prev_speaker] = []

                    if prev_speaker not in speakers_order:  # if the speaker is not in the order then add him
                        speakers_order.append(prev_speaker)

                elif prev_speaker != '':  # if we have a speaker then add the text to his name
                    split_txt = split_paragrph(par)
                    filtered_text = []
                    for sent in split_txt:
                        filtered = clean_text(sent)
                        if filtered != '':
                            filtered_text.append(filtered)

                    all_tokens = tokenize(filtered_text)
                    if len(all_tokens) == 0:
                        continue

                    for token in all_tokens:
                        combine_tokens = ''
                        for word1 in token:
                            combine_tokens += str(word1) + ' '
                        speaker_text[prev_speaker].append(combine_tokens.strip())

            for speaker in speakers_order:
                for text in speaker_text[speaker]:
                    jsonl_data.append({
                        'protocol_name': protocol_name,
                        'knesset_number': knesset_number,
                        'protocol_type': protocol_type,
                        'protocol_number': protocol_int,
                        'speaker_name': speaker,
                        'sentence_text': text
                    })

        with open(output_path, 'w', encoding='utf-8') as jsonl_file:
            for data_item in jsonl_data:  # change back
                # Convert the dictionary to json lines
                json_line = json.dumps(data_item, ensure_ascii=False)

                # Write the json line to the file
                jsonl_file.write(json_line + '\n')

    except Exception as e:
        # Handle any exception
        print(f"An error occurred in main: {e}")
